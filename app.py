import streamlit as st
from pathlib import Path
from io import BytesIO
from datetime import datetime
import re, random, html

from pypdf import PdfReader
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak

st.set_page_config(page_title="PaperCraft Pro", page_icon="📘", layout="wide", initial_sidebar_state="expanded")

st.markdown("""
<style>
#MainMenu {visibility:hidden} footer {visibility:hidden}
.stApp {background:#f7f8fc}
.block-container {padding-top:1.4rem; max-width:1450px}
.brand {font-size:2.25rem;font-weight:850;letter-spacing:-1px;margin-bottom:.1rem}
.brand span{color:#f04f23}
.sub {color:#697386;font-size:1.05rem;margin-bottom:1.6rem}
.hero {background:linear-gradient(115deg,#ef4d22,#ff7a18 55%,#ffb000);border-radius:26px;padding:42px;color:white;margin:12px 0 28px;box-shadow:0 18px 45px rgba(240,79,35,.22)}
.hero h1{font-size:3.1rem;margin:0 0 8px;font-weight:850}
.hero p{font-size:1.15rem;max-width:760px;opacity:.96}
.toolcard{background:white;border:1px solid #eceef4;border-radius:20px;padding:22px;height:100%;box-shadow:0 5px 18px rgba(30,40,70,.05)}
.toolcard h3{margin-top:0}
.smallcard{background:white;border:1px solid #eceef4;border-radius:16px;padding:17px}
div.stButton>button, div.stDownloadButton>button {border-radius:12px;font-weight:700;min-height:44px;border:0}
div.stButton>button[kind="primary"]{background:#ef4d22}
[data-testid="stSidebar"]{background:#171a21}
[data-testid="stSidebar"] *{color:#f3f4f6}
.section-title{font-size:1.55rem;font-weight:800;margin:15px 0}
.badge{display:inline-block;background:#fff1eb;color:#d53f18;padding:5px 11px;border-radius:20px;font-size:.82rem;font-weight:700;margin-right:6px}
hr{border-color:#eceef4}
</style>
""", unsafe_allow_html=True)

# ---------- DATA ----------
def init():
    defaults = {
        "material_text":"", "material_names":[], "questions":[], "paper_title":"Final Examination",
        "school":"", "subject":"", "grade":"", "duration":"2 Hours", "total_marks":100,
        "instructions":"Answer all questions. Read every question carefully before answering.",
        "logo":None, "page":"Dashboard", "history":[]
    }
    for k,v in defaults.items():
        if k not in st.session_state: st.session_state[k]=v
init()

def extract_file(f):
    name=f.name.lower()
    if name.endswith(".pdf"):
        return "\n".join((p.extract_text() or "") for p in PdfReader(f).pages)
    if name.endswith(".docx"):
        d=Document(f); return "\n".join(p.text for p in d.paragraphs)
    if name.endswith(".txt"):
        return f.read().decode("utf-8", errors="ignore")
    return ""

def split_sentences(text):
    text=re.sub(r"\s+"," ",text).strip()
    return [x.strip() for x in re.split(r'(?<=[.!?])\s+',text) if len(x.split())>=7]

def keywords(text):
    words=re.findall(r"[A-Za-z][A-Za-z\-]{3,}",text.lower())
    stop=set("""this that these those with from have has had were where which while about into also than then they them their there what when will would could should such more most each every using used use only very some other over under between after before""".split())
    freq={}
    for w in words:
        if w not in stop: freq[w]=freq.get(w,0)+1
    return [w for w,_ in sorted(freq.items(),key=lambda x:x[1],reverse=True)]

def difficulty_label(i):
    return ["Easy","Medium","Hard"][i%3]

def generate_questions(text, mcq, short, longq, tf, difficulty, bloom):
    s=split_sentences(text)
    if len(s)<2: raise ValueError("Please upload material containing more readable educational text.")
    random.shuffle(s)
    keys=keywords(text)
    qs=[]; idx=0

    def source():
        nonlocal idx
        x=s[idx%len(s)]; idx+=1
        return x

    for _ in range(mcq):
        sentence=source()
        words=re.findall(r"[A-Za-z][A-Za-z\-]{4,}",sentence)
        answer=max(words,key=len).title() if words else "Concept"
        pool=[k.title() for k in keys if k.title()!=answer]
        random.shuffle(pool)
        options=[answer]+pool[:3]
        while len(options)<4: options.append("None of the above")
        random.shuffle(options)
        qs.append({"type":"Multiple Choice","question":f"Based on the study material, which option best relates to this statement?<br><br><i>{html.escape(sentence)}</i>",
                   "options":options,"answer":answer,"marks":1,"difficulty":difficulty or difficulty_label(len(qs)),"bloom":bloom})

    for _ in range(tf):
        sentence=source()
        truth=random.choice([True,False])
        qtext=sentence if truth else "The following statement is incorrect according to the material: "+sentence
        qs.append({"type":"True / False","question":qtext,"options":["True","False"],"answer":"True" if truth else "False","marks":1,"difficulty":difficulty or "Easy","bloom":bloom})

    for _ in range(short):
        sentence=source()
        qs.append({"type":"Short Answer","question":"Explain briefly the following concept or idea from the teaching material:<br><br>"+html.escape(sentence),
                   "options":[],"answer":sentence,"marks":3,"difficulty":difficulty or "Medium","bloom":bloom})

    for _ in range(longq):
        sentence=source()
        qs.append({"type":"Long Answer","question":"Discuss the following topic in detail. Include the main concepts, significance, examples, and implications:<br><br>"+html.escape(sentence),
                   "options":[],"answer":sentence,"marks":8,"difficulty":difficulty or "Hard","bloom":bloom})
    return qs

# ---------- EXPORT ----------
def clean_html(s): return re.sub("<[^>]+>","",s)

def build_docx(answer_key=False):
    d=Document()
    sec=d.sections[0]; sec.top_margin=15*mm; sec.bottom_margin=15*mm
    if st.session_state.school: d.add_heading(st.session_state.school,0)
    d.add_heading(st.session_state.paper_title,1)
    d.add_paragraph(f"Subject: {st.session_state.subject} | Class: {st.session_state.grade}")
    d.add_paragraph(f"Duration: {st.session_state.duration} | Total Marks: {st.session_state.total_marks}")
    d.add_paragraph("Instructions: "+st.session_state.instructions)
    d.add_paragraph("")
    current=None
    for i,q in enumerate(st.session_state.questions,1):
        if q["type"]!=current:
            current=q["type"]; d.add_heading(current,2)
        d.add_paragraph(f"{i}. {clean_html(q['question'])}  [{q['marks']} marks]")
        for j,o in enumerate(q["options"]): d.add_paragraph(f"    {chr(65+j)}. {o}")
        if answer_key: d.add_paragraph("ANSWER: "+q["answer"])
    out=BytesIO(); d.save(out); out.seek(0); return out

def build_pdf(answer_key=False):
    out=BytesIO()
    doc=SimpleDocTemplate(out,pagesize=A4,rightMargin=18*mm,leftMargin=18*mm,topMargin=16*mm,bottomMargin=16*mm)
    ss=getSampleStyleSheet()
    title=ParagraphStyle("T",parent=ss["Title"],textColor=colors.HexColor("#EF4D22"),alignment=1,fontSize=22,spaceAfter=8)
    h=ParagraphStyle("H",parent=ss["Heading2"],textColor=colors.HexColor("#EF4D22"),spaceBefore=13)
    story=[]
    if st.session_state.school: story.append(Paragraph(html.escape(st.session_state.school),ss["Heading2"]))
    story += [Paragraph(html.escape(st.session_state.paper_title),title),
              Paragraph(html.escape(f"Subject: {st.session_state.subject} | Class: {st.session_state.grade}"),ss["Normal"]),
              Paragraph(html.escape(f"Duration: {st.session_state.duration} | Total Marks: {st.session_state.total_marks}"),ss["Normal"]),
              Spacer(1,7),Paragraph("<b>Instructions:</b> "+html.escape(st.session_state.instructions),ss["Normal"]),Spacer(1,7)]
    current=None
    for i,q in enumerate(st.session_state.questions,1):
        if q["type"]!=current:
            current=q["type"]; story.append(Paragraph(current,h))
        story.append(Paragraph(f"<b>{i}.</b> {q['question']} &nbsp; <b>[{q['marks']} marks]</b>",ss["BodyText"]))
        for j,o in enumerate(q["options"]):
            story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;{chr(65+j)}. {html.escape(o)}",ss["BodyText"]))
        if answer_key:
            story.append(Paragraph(f"<font color='#EF4D22'><b>Answer:</b> {html.escape(q['answer'])}</font>",ss["BodyText"]))
        story.append(Spacer(1,5))
    doc.build(story); out.seek(0); return out

# ---------- SIDEBAR ----------
with st.sidebar:
    st.markdown("## 📘 PaperCraft Pro")
    st.caption("AI-style exam workspace")
    pages=["Dashboard","Upload Material","Question Studio","Paper Designer","Export Center"]
    icons=["🏠","📤","✨","🎨","⬇️"]
    for p,ic in zip(pages,icons):
        if st.button(f"{ic}  {p}",use_container_width=True):
            st.session_state.page=p
    st.divider()
    st.caption("Professional exam creation platform")
    st.caption("PDF • Word • Answer Keys")

# ---------- DASHBOARD ----------
page=st.session_state.page
if page=="Dashboard":
    st.markdown('<div class="brand">PaperCraft <span>Pro</span></div>',unsafe_allow_html=True)
    st.markdown('<div class="sub">Create professional examination papers from your teaching material.</div>',unsafe_allow_html=True)
    st.markdown("""<div class="hero"><h1>Exam papers. Reimagined. ⚡</h1>
    <p>Upload your study material, generate structured questions, design your paper, and export professional PDF or Word documents — all from one workspace.</p></div>""",unsafe_allow_html=True)
    c1,c2,c3,c4=st.columns(4)
    c1.metric("📄 Materials",len(st.session_state.material_names))
    c2.metric("❓ Questions",len(st.session_state.questions))
    c3.metric("🧮 Current Marks",sum(q["marks"] for q in st.session_state.questions))
    c4.metric("📦 Exports Ready",2 if st.session_state.questions else 0)
    st.markdown("### Your tools")
    a,b,c=st.columns(3)
    with a: st.markdown('<div class="toolcard"><h3>📤 Upload Material</h3><p>Build your knowledge base from PDF, Word and text files.</p></div>',unsafe_allow_html=True)
    with b: st.markdown('<div class="toolcard"><h3>✨ Question Studio</h3><p>Create MCQs, true/false, short and long questions.</p></div>',unsafe_allow_html=True)
    with c: st.markdown('<div class="toolcard"><h3>🎨 Export Center</h3><p>Download polished exam papers and answer keys.</p></div>',unsafe_allow_html=True)

elif page=="Upload Material":
    st.markdown("## 📤 Upload Material")
    st.caption("Create the knowledge base your paper will be generated from.")
    files=st.file_uploader("Drop your teaching files here",type=["pdf","docx","txt"],accept_multiple_files=True)
    if files and st.button("Process Material",type="primary"):
        text=[]; names=[]
        for f in files:
            try:
                text.append(extract_file(f)); names.append(f.name)
            except Exception as e: st.warning(f"{f.name}: {e}")
        st.session_state.material_text="\n\n".join(text)
        st.session_state.material_names=names
        st.success(f"Knowledge base created from {len(names)} file(s): {len(st.session_state.material_text):,} characters.")
    if st.session_state.material_text:
        st.markdown("### Knowledge Base")
        st.markdown(" ".join(f'<span class="badge">📄 {html.escape(x)}</span>' for x in st.session_state.material_names),unsafe_allow_html=True)
        with st.expander("Preview extracted content"): st.write(st.session_state.material_text[:7000])

elif page=="Question Studio":
    st.markdown("## ✨ Question Studio")
    if not st.session_state.material_text:
        st.warning("Upload teaching material first.")
        st.stop()
    col1,col2=st.columns([1,1])
    with col1:
        st.markdown("### Question formats")
        mcq=st.number_input("🔘 Multiple Choice",0,50,5)
        tf=st.number_input("✓ True / False",0,30,2)
        short=st.number_input("✍️ Short Answer",0,30,5)
        longq=st.number_input("📝 Long Answer",0,20,3)
    with col2:
        st.markdown("### Academic controls")
        difficulty=st.selectbox("Difficulty",["Mixed","Easy","Medium","Hard"])
        bloom=st.selectbox("Bloom's Taxonomy",["Understand","Apply","Analyze","Evaluate","Create"])
        regenerate=st.checkbox("Replace existing questions",True)
        st.info("The generator uses the uploaded material as its knowledge source.")
    if st.button("🚀 Generate Question Set",type="primary",use_container_width=True):
        d="" if difficulty=="Mixed" else difficulty
        new=generate_questions(st.session_state.material_text,int(mcq),int(short),int(longq),int(tf),d,bloom)
        st.session_state.questions=new if regenerate else st.session_state.questions+new
        st.success(f"Generated {len(new)} questions.")
    if st.session_state.questions:
        st.markdown("### Generated Questions")
        for i,q in enumerate(st.session_state.questions):
            with st.expander(f"{i+1}. {q['type']} • {q['difficulty']} • {q['marks']} marks"):
                st.markdown(q["question"],unsafe_allow_html=True)
                if q["options"]:
                    for o in q["options"]: st.write("• "+o)
                q["marks"]=st.number_input("Marks",1,100,q["marks"],key=f"marks{i}")
                if st.button("🗑 Remove question",key=f"del{i}"):
                    st.session_state.questions.pop(i); st.rerun()

elif page=="Paper Designer":
    st.markdown("## 🎨 Paper Designer")
    st.caption("Customize the identity and structure of your examination paper.")
    a,b=st.columns(2)
    with a:
        st.session_state.school=st.text_input("School / University",st.session_state.school)
        st.session_state.paper_title=st.text_input("Paper Title",st.session_state.paper_title)
        st.session_state.subject=st.text_input("Subject",st.session_state.subject)
        st.session_state.grade=st.text_input("Class / Grade",st.session_state.grade)
    with b:
        st.session_state.duration=st.text_input("Duration",st.session_state.duration)
        st.session_state.total_marks=st.number_input("Official Total Marks",1,1000,st.session_state.total_marks)
        st.session_state.instructions=st.text_area("Instructions",st.session_state.instructions,height=130)
        st.session_state.logo=st.file_uploader("Optional School Logo",type=["png","jpg","jpeg"])
    st.markdown("### Paper Summary")
    st.info(f"{len(st.session_state.questions)} questions • Generated marks: {sum(q['marks'] for q in st.session_state.questions)} • Target marks: {st.session_state.total_marks}")
    st.markdown("### Live Preview")
    st.markdown(f'<div class="toolcard"><h2>{html.escape(st.session_state.school)}</h2><h1>{html.escape(st.session_state.paper_title)}</h1><p><b>{html.escape(st.session_state.subject)}</b> &nbsp; | &nbsp; {html.escape(st.session_state.grade)} &nbsp; | &nbsp; {html.escape(st.session_state.duration)}</p><hr><p>{html.escape(st.session_state.instructions)}</p></div>',unsafe_allow_html=True)

elif page=="Export Center":
    st.markdown("## ⬇️ Export Center")
    if not st.session_state.questions:
        st.warning("Generate questions before exporting.")
        st.stop()
    st.markdown('<div class="hero"><h1>Your paper is ready 🎉</h1><p>Download professional documents for printing, sharing, or editing.</p></div>',unsafe_allow_html=True)
    a,b,c,d=st.columns(4)
    with a:
        st.download_button("📕 Exam PDF",build_pdf(False),f"{st.session_state.paper_title}.pdf","application/pdf",use_container_width=True)
    with b:
        st.download_button("📘 Exam Word",build_docx(False),f"{st.session_state.paper_title}.docx","application/vnd.openxmlformats-officedocument.wordprocessingml.document",use_container_width=True)
    with c:
        st.download_button("🔐 Answer Key PDF",build_pdf(True),f"{st.session_state.paper_title}_Answer_Key.pdf","application/pdf",use_container_width=True)
    with d:
        st.download_button("📝 Answer Key Word",build_docx(True),f"{st.session_state.paper_title}_Answer_Key.docx","application/vnd.openxmlformats-officedocument.wordprocessingml.document",use_container_width=True)
    st.success("All exports are generated directly from your current PaperCraft Pro workspace.")
