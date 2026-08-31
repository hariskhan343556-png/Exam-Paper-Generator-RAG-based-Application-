import os
import json
import hashlib
import tempfile
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple
import streamlit as st
import PyPDF2
import docx
import pandas as pd
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from sentence_transformers import SentenceTransformer
import chromadb
from chromadb.utils import embedding_functions
import torch
import numpy as np
from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM
import re
import nltk
from nltk.tokenize import sent_tokenize
from nltk.corpus import stopwords
from collections import Counter
import random
from io import BytesIO

nltk.download('punkt', quiet=True)
nltk.download('stopwords', quiet=True)

class DocumentProcessor:
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
    
    def extract_text(self, file_content: bytes, file_name: str) -> str:
        file_extension = os.path.splitext(file_name)[1].lower()
        
        if file_extension == '.pdf':
            return self._extract_pdf(file_content)
        elif file_extension == '.docx':
            return self._extract_docx(file_content)
        elif file_extension == '.txt':
            return self._extract_txt(file_content)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _extract_pdf(self, file_content: bytes) -> str:
        text = ""
        with tempfile.NamedTemporaryFile(delete=True, suffix='.pdf') as tmp_file:
            tmp_file.write(file_content)
            tmp_file.flush()
            
            with open(tmp_file.name, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        return text
    
    def _extract_docx(self, file_content: bytes) -> str:
        with tempfile.NamedTemporaryFile(delete=True, suffix='.docx') as tmp_file:
            tmp_file.write(file_content)
            tmp_file.flush()
            
            doc = docx.Document(tmp_file.name)
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text
    
    def _extract_txt(self, file_content: bytes) -> str:
        return file_content.decode('utf-8', errors='ignore')
    
    def chunk_text(self, text: str, chunk_size: int = 500, overlap: int = 100) -> List[str]:
        sentences = sent_tokenize(text)
        chunks = []
        current_chunk = []
        current_size = 0
        
        for sentence in sentences:
            sentence_len = len(sentence.split())
            
            if current_size + sentence_len > chunk_size and current_chunk:
                chunks.append(" ".join(current_chunk))
                
                overlap_sentences = current_chunk[-min(len(current_chunk), overlap//10):]
                current_chunk = overlap_sentences
                current_size = sum(len(s.split()) for s in current_chunk)
            
            current_chunk.append(sentence)
            current_size += sentence_len
        
        if current_chunk:
            chunks.append(" ".join(current_chunk))
        
        return chunks

class VectorStore:
    def __init__(self, collection_name: str = "exam_materials"):
        self.model = SentenceTransformer('all-MiniLM-L6-v2')
        self.client = chromadb.PersistentClient(path="./chroma_db")
        self.collection_name = collection_name
        
        try:
            self.collection = self.client.get_collection(name=collection_name)
        except ValueError:
            self.collection = self.client.create_collection(
                name=collection_name,
                embedding_function=embedding_functions.SentenceTransformerEmbeddingFunction(
                    model_name='all-MiniLM-L6-v2'
                )
            )
    
    def add_documents(self, chunks: List[str], metadata: Optional[Dict] = None):
        if not chunks:
            return
        
        ids = [hashlib.md5(chunk.encode()).hexdigest() for chunk in chunks]
        metadatas = [metadata or {} for _ in chunks]
        
        existing_ids = set(self.collection.get()['ids']) if self.collection.count() > 0 else set()
        new_chunks = []
        new_ids = []
        new_metadatas = []
        
        for chunk, chunk_id, meta in zip(chunks, ids, metadatas):
            if chunk_id not in existing_ids:
                new_chunks.append(chunk)
                new_ids.append(chunk_id)
                new_metadatas.append(meta)
        
        if new_chunks:
            self.collection.add(
                documents=new_chunks,
                ids=new_ids,
                metadatas=new_metadatas
            )
    
    def query(self, query_text: str, n_results: int = 5) -> List[str]:
        if self.collection.count() == 0:
            return []
        
        results = self.collection.query(
            query_texts=[query_text],
            n_results=min(n_results, self.collection.count())
        )
        
        return results['documents'][0] if results['documents'] else []
    
    def clear(self):
        self.client.delete_collection(self.collection_name)
        self.collection = self.client.create_collection(
            name=self.collection_name,
            embedding_function=embedding_functions.SentenceTransformerEmbeddingFunction(
                model_name='all-MiniLM-L6-v2'
            )
        )

class QuestionGenerator:
    def __init__(self):
        self.model_name = "t5-small"
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        
        self.qg_pipeline = pipeline(
            "text2text-generation",
            model="valhalla/t5-small-qg-hl",
            device=0 if self.device == "cuda" else -1
        )
        
        self.stop_words = set(stopwords.words('english'))
    
    def generate_questions(self, context: str, num_questions: int = 5, 
                          question_type: str = "mcq") -> List[Dict[str, Any]]:
        if question_type == "mcq":
            return self._generate_mcqs(context, num_questions)
        elif question_type == "short":
            return self._generate_short_questions(context, num_questions)
        elif question_type == "long":
            return self._generate_long_questions(context, num_questions)
        else:
            return self._generate_mcqs(context, num_questions)
    
    def _generate_mcqs(self, context: str, num_questions: int) -> List[Dict[str, Any]]:
        questions = []
        sentences = sent_tokenize(context)
        
        if not sentences:
            return questions
        
        selected_sentences = random.sample(sentences, min(num_questions * 2, len(sentences)))
        
        for sentence in selected_sentences[:num_questions]:
            if len(sentence.split()) < 5:
                continue
            
            try:
                question_text = f"question: {sentence} context: {context[:200]}"
                result = self.qg_pipeline(question_text, max_length=64, num_return_sequences=1)[0]['generated_text']
                
                if result:
                    options = self._generate_options(context, sentence, 4)
                    
                    question = {
                        'question': result,
                        'options': options,
                        'answer': sentence,
                        'type': 'mcq'
                    }
                    questions.append(question)
            except Exception as e:
                continue
            
            if len(questions) >= num_questions:
                break
        
        while len(questions) < num_questions:
            fallback_sentence = random.choice(sentences) if sentences else "Sample question based on the material."
            question = {
                'question': f"Based on the material, what is the main idea of: {fallback_sentence[:100]}?",
                'options': self._generate_options(context, fallback_sentence, 4),
                'answer': fallback_sentence,
                'type': 'mcq'
            }
            questions.append(question)
        
        return questions[:num_questions]
    
    def _generate_options(self, context: str, correct_answer: str, num_options: int) -> List[str]:
        sentences = sent_tokenize(context)
        wrong_answers = [s for s in sentences if s != correct_answer and len(s.split()) > 3]
        
        options = [correct_answer]
        
        if len(wrong_answers) >= num_options - 1:
            selected_wrong = random.sample(wrong_answers, num_options - 1)
            options.extend(selected_wrong)
        else:
            options.extend(wrong_answers[:num_options - 1])
            while len(options) < num_options:
                filler = f"Alternative concept related to {random.choice(options)[:20]}"
                if filler not in options:
                    options.append(filler)
        
        random.shuffle(options)
        return options
    
    def _generate_short_questions(self, context: str, num_questions: int) -> List[Dict[str, Any]]:
        questions = []
        sentences = sent_tokenize(context)
        
        if not sentences:
            return questions
        
        selected_sentences = random.sample(sentences, min(num_questions * 2, len(sentences)))
        
        for sentence in selected_sentences[:num_questions]:
            if len(sentence.split()) < 5:
                continue
            
            try:
                question_text = f"question: {sentence} context: {context[:200]}"
                result = self.qg_pipeline(question_text, max_length=64, num_return_sequences=1)[0]['generated_text']
                
                if result:
                    question = {
                        'question': result,
                        'answer': sentence,
                        'type': 'short',
                        'expected_length': 50
                    }
                    questions.append(question)
            except Exception as e:
                continue
            
            if len(questions) >= num_questions:
                break
        
        while len(questions) < num_questions:
            fallback_sentence = random.choice(sentences) if sentences else "Key concept from the material."
            question = {
                'question': f"Explain the following concept: {fallback_sentence[:100]}",
                'answer': fallback_sentence,
                'type': 'short',
                'expected_length': 50
            }
            questions.append(question)
        
        return questions[:num_questions]
    
    def _generate_long_questions(self, context: str, num_questions: int) -> List[Dict[str, Any]]:
        questions = []
        sentences = sent_tokenize(context)
        
        if not sentences:
            return questions
        
        chunk_size = max(3, len(sentences) // max(1, num_questions))
        chunks = [sentences[i:i+chunk_size] for i in range(0, len(sentences), chunk_size)]
        
        for chunk in chunks[:num_questions]:
            if len(chunk) < 2:
                continue
            
            chunk_text = " ".join(chunk)
            
            try:
                question_text = f"question: {chunk_text[:150]} context: {context[:300]}"
                result = self.qg_pipeline(question_text, max_length=128, num_return_sequences=1)[0]['generated_text']
                
                if result:
                    question = {
                        'question': result,
                        'answer': chunk_text,
                        'type': 'long',
                        'expected_length': 200
                    }
                    questions.append(question)
            except Exception as e:
                continue
            
            if len(questions) >= num_questions:
                break
        
        while len(questions) < num_questions:
            fallback_chunk = " ".join(random.sample(sentences, min(5, len(sentences)))) if sentences else "Comprehensive topic from the material."
            question = {
                'question': f"Analyze and discuss: {fallback_chunk[:150]}",
                'answer': fallback_chunk,
                'type': 'long',
                'expected_length': 200
            }
            questions.append(question)
        
        return questions[:num_questions]

class ExamPaperGenerator:
    def __init__(self):
        self.pdf_template = SimpleDocTemplate
        self.word_template = docx.Document
    
    def generate_pdf(self, questions: List[Dict[str, Any]], title: str = "Exam Paper") -> BytesIO:
        buffer = BytesIO()
        doc = self.pdf_template(buffer, pagesize=letter)
        
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1a1a2e'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#16213e'),
            spaceAfter=12
        )
        
        question_style = ParagraphStyle(
            'QuestionStyle',
            parent=styles['Normal'],
            fontSize=12,
            spaceAfter=10,
            alignment=TA_LEFT
        )
        
        option_style = ParagraphStyle(
            'OptionStyle',
            parent=styles['Normal'],
            fontSize=11,
            leftIndent=20,
            spaceAfter=6
        )
        
        story = []
        
        story.append(Paragraph(f"{title}", title_style))
        story.append(Paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        section_headers = {
            'mcq': 'Multiple Choice Questions',
            'short': 'Short Answer Questions',
            'long': 'Long Answer Questions'
        }
        
        for q_type in ['mcq', 'short', 'long']:
            type_questions = [q for q in questions if q.get('type') == q_type]
            
            if not type_questions:
                continue
            
            story.append(PageBreak())
            story.append(Paragraph(section_headers.get(q_type, ''), heading_style))
            story.append(Spacer(1, 0.1*inch))
            
            for idx, q in enumerate(type_questions, 1):
                question_text = f"{idx}. {q['question']}"
                story.append(Paragraph(question_text, question_style))
                
                if q.get('type') == 'mcq':
                    options = q.get('options', [])
                    for opt_idx, option in enumerate(options, 1):
                        option_text = f"{chr(64 + opt_idx)}. {option}"
                        story.append(Paragraph(option_text, option_style))
                    story.append(Spacer(1, 0.05*inch))
                elif q.get('type') == 'short':
                    story.append(Paragraph("Answer: (3-5 lines)", styles['Normal']))
                    story.append(Spacer(1, 0.3*inch))
                    story.append(Paragraph("-" * 60, styles['Normal']))
                    story.append(Spacer(1, 0.05*inch))
                elif q.get('type') == 'long':
                    story.append(Paragraph("Answer: (10-15 lines)", styles['Normal']))
                    story.append(Spacer(1, 0.6*inch))
                    story.append(Paragraph("-" * 60, styles['Normal']))
                    story.append(Spacer(1, 0.05*inch))
        
        doc.build(story)
        buffer.seek(0)
        return buffer
    
    def generate_word(self, questions: List[Dict[str, Any]], title: str = "Exam Paper") -> BytesIO:
        buffer = BytesIO()
        doc = docx.Document()
        
        title_paragraph = doc.add_heading(title, 0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date_paragraph = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y')}")
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_page_break()
        
        section_headers = {
            'mcq': 'Multiple Choice Questions',
            'short': 'Short Answer Questions',
            'long': 'Long Answer Questions'
        }
        
        for q_type in ['mcq', 'short', 'long']:
            type_questions = [q for q in questions if q.get('type') == q_type]
            
            if not type_questions:
                continue
            
            doc.add_page_break()
            doc.add_heading(section_headers.get(q_type, ''), 1)
            
            for idx, q in enumerate(type_questions, 1):
                p = doc.add_paragraph()
                run = p.add_run(f"{idx}. {q['question']}")
                run.bold = True
                
                if q.get('type') == 'mcq':
                    options = q.get('options', [])
                    for opt_idx, option in enumerate(options, 1):
                        p = doc.add_paragraph()
                        p.add_run(f"{chr(64 + opt_idx)}. {option}")
                        p.paragraph_format.left_indent = Inches(0.3)
                
                elif q.get('type') == 'short':
                    doc.add_paragraph("Answer: (3-5 lines)")
                    doc.add_paragraph("_" * 60)
                    doc.add_paragraph()
                
                elif q.get('type') == 'long':
                    doc.add_paragraph("Answer: (10-15 lines)")
                    doc.add_paragraph("_" * 60)
                    doc.add_paragraph()
        
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    def generate_paper(self, questions: List[Dict[str, Any]], format: str = "pdf", 
                      title: str = "Exam Paper") -> BytesIO:
        if format.lower() == "pdf":
            return self.generate_pdf(questions, title)
        elif format.lower() == "word":
            return self.generate_word(questions, title)
        else:
            raise ValueError(f"Unsupported format: {format}")

class ExamPaperApp:
    def __init__(self):
        self.document_processor = DocumentProcessor()
        self.vector_store = VectorStore()
        self.question_generator = QuestionGenerator()
        self.exam_generator = ExamPaperGenerator()
        self.uploaded_files = []
        
    def run(self):
        st.set_page_config(
            page_title="Exam Paper Generator",
            page_icon="📝",
            layout="wide"
        )
        
        st.title("📝 Exam Paper Generator - RAG Application")
        st.markdown("---")
        
        with st.sidebar:
            st.header("📚 Upload Teaching Material")
            
            uploaded_files = st.file_uploader(
                "Upload PDF, DOCX, or TXT files",
                type=['pdf', 'docx', 'txt'],
                accept_multiple_files=True
            )
            
            if uploaded_files:
                if st.button("Process Documents", type="primary"):
                    with st.spinner("Processing documents..."):
                        for file in uploaded_files:
                            try:
                                text = self.document_processor.extract_text(
                                    file.getvalue(), 
                                    file.name
                                )
                                
                                chunks = self.document_processor.chunk_text(text)
                                
                                metadata = {
                                    'filename': file.name,
                                    'upload_time': datetime.now().isoformat()
                                }
                                
                                self.vector_store.add_documents(chunks, metadata)
                                
                                self.uploaded_files.append({
                                    'name': file.name,
                                    'chunks': len(chunks)
                                })
                            except Exception as e:
                                st.error(f"Error processing {file.name}: {str(e)}")
                        
                        st.success(f"Processed {len(uploaded_files)} files successfully!")
            
            st.markdown("---")
            
            st.header("⚙️ Exam Configuration")
            
            num_questions = st.slider(
                "Number of questions per section",
                min_value=1,
                max_value=20,
                value=5
            )
            
            include_mcq = st.checkbox("Multiple Choice Questions", value=True)
            include_short = st.checkbox("Short Answer Questions", value=True)
            include_long = st.checkbox("Long Answer Questions", value=True)
            
            difficulty = st.selectbox(
                "Difficulty Level",
                ["Easy", "Medium", "Hard"],
                index=1
            )
            
            output_format = st.radio(
                "Output Format",
                ["PDF", "Word"],
                index=0
            )
            
            st.markdown("---")
            
            if st.button("📄 Generate Exam Paper", type="primary", use_container_width=True):
                self._generate_paper(
                    num_questions,
                    include_mcq,
                    include_short,
                    include_long,
                    difficulty,
                    output_format
                )
        
        col1, col2 = st.columns([2, 1])
        
        with col1:
            st.header("📖 Current Material")
            
            if self.vector_store.collection.count() > 0:
                st.info(f"Total chunks in database: {self.vector_store.collection.count()}")
                
                for file_info in self.uploaded_files:
                    st.markdown(f"- ✅ **{file_info['name']}** ({file_info['chunks']} chunks)")
                
                if st.button("Clear All Materials"):
                    self.vector_store.clear()
                    self.uploaded_files = []
                    st.success("All materials cleared!")
            else:
                st.info("No materials uploaded yet. Please upload teaching material using the sidebar.")
        
        with col2:
            st.header("📊 Statistics")
            
            if self.vector_store.collection.count() > 0:
                st.metric("Total Documents", len(self.uploaded_files))
                st.metric("Total Chunks", self.vector_store.collection.count())
            else:
                st.metric("Total Documents", "0")
                st.metric("Total Chunks", "0")
    
    def _generate_paper(self, num_questions, include_mcq, include_short, 
                        include_long, difficulty, output_format):
        
        if self.vector_store.collection.count() == 0:
            st.error("Please upload teaching material first!")
            return
        
        with st.spinner("Generating exam paper..."):
            try:
                retrieved_chunks = []
                for topic in ["important concepts", "key definitions", "main principles"]:
                    chunks = self.vector_store.query(topic, n_results=5)
                    retrieved_chunks.extend(chunks)
                
                if not retrieved_chunks:
                    st.error("Could not retrieve relevant content. Please upload more material.")
                    return
                
                context = " ".join(retrieved_chunks)
                
                all_questions = []
                
                if include_mcq:
                    mcq_questions = self.question_generator.generate_questions(
                        context, 
                        num_questions, 
                        "mcq"
                    )
                    all_questions.extend(mcq_questions)
                
                if include_short:
                    short_questions = self.question_generator.generate_questions(
                        context, 
                        num_questions, 
                        "short"
                    )
                    all_questions.extend(short_questions)
                
                if include_long:
                    long_questions = self.question_generator.generate_questions(
                        context, 
                        max(1, num_questions // 2), 
                        "long"
                    )
                    all_questions.extend(long_questions)
                
                if not all_questions:
                    st.error("Failed to generate questions. Please try again.")
                    return
                
                title = f"Exam Paper - {datetime.now().strftime('%Y-%m-%d')}"
                
                if output_format == "PDF":
                    buffer = self.exam_generator.generate_pdf(all_questions, title)
                    mime_type = "application/pdf"
                    file_extension = ".pdf"
                else:
                    buffer = self.exam_generator.generate_word(all_questions, title)
                    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    file_extension = ".docx"
                
                st.success(f"✅ Exam paper generated successfully with {len(all_questions)} questions!")
                
                st.download_button(
                    label=f"Download {output_format}",
                    data=buffer,
                    file_name=f"exam_paper_{datetime.now().strftime('%Y%m%d_%H%M%S')}{file_extension}",
                    mime=mime_type,
                    use_container_width=True
                )
                
                with st.expander("Preview Questions"):
                    for idx, q in enumerate(all_questions, 1):
                        st.markdown(f"**{idx}. {q['question']}**")
                        if q.get('type') == 'mcq':
                            options = q.get('options', [])
                            for opt_idx, opt in enumerate(options, 1):
                                st.markdown(f"&nbsp;&nbsp;&nbsp;&nbsp;{chr(64+opt_idx)}. {opt}")
                        st.markdown("---")
                
            except Exception as e:
                st.error(f"Error generating exam paper: {str(e)}")

def main():
    app = ExamPaperApp()
    app.run()

if __name__ == "__main__":
    main()