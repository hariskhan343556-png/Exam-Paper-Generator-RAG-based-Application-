import os
import json
import hashlib
import tempfile
import re
import random
import string
from datetime import datetime
from typing import List, Dict, Any, Optional, Tuple, Union
from collections import Counter, defaultdict
import streamlit as st
import PyPDF2
import docx
import pandas as pd
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, cm
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY, TA_RIGHT
from reportlab.pdfgen import canvas
from sentence_transformers import SentenceTransformer
import chromadb
from chromadb.utils import embedding_functions
import torch
import numpy as np
from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM, AutoModelForQuestionAnswering
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from nltk.probability import FreqDist
from nltk.stem import WordNetLemmatizer
import spacy
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import networkx as nx
import pickle
from io import BytesIO
import base64
import hashlib
import time
from pathlib import Path

nltk.download('punkt', quiet=True)
nltk.download('stopwords', quiet=True)
nltk.download('wordnet', quiet=True)
nltk.download('averaged_perceptron_tagger', quiet=True)

class AdvancedTextProcessor:
    def __init__(self):
        self.lemmatizer = WordNetLemmatizer()
        self.stop_words = set(stopwords.words('english'))
        try:
            self.nlp = spacy.load('en_core_web_sm')
        except:
            import subprocess
            subprocess.run(['python', '-m', 'spacy', 'download', 'en_core_web_sm'])
            self.nlp = spacy.load('en_core_web_sm')
    
    def extract_key_phrases(self, text: str, top_n: int = 20) -> List[str]:
        doc = self.nlp(text)
        
        phrases = []
        for chunk in doc.noun_chunks:
            if len(chunk.text.split()) > 1:
                phrases.append(chunk.text.lower())
        
        for ent in doc.ents:
            if ent.label_ in ['ORG', 'PERSON', 'GPE', 'PRODUCT']:
                phrases.append(ent.text.lower())
        
        freq_dist = Counter(phrases)
        return [phrase for phrase, _ in freq_dist.most_common(top_n)]
    
    def extract_keywords(self, text: str, top_n: int = 30) -> List[str]:
        doc = self.nlp(text)
        keywords = []
        
        for token in doc:
            if token.is_stop or token.is_punct or token.is_space:
                continue
            if token.pos_ in ['NOUN', 'PROPN', 'ADJ', 'VERB']:
                lemmatized = self.lemmatizer.lemmatize(token.text.lower())
                keywords.append(lemmatized)
        
        freq_dist = Counter(keywords)
        return [word for word, _ in freq_dist.most_common(top_n)]
    
    def get_sentence_importance(self, sentences: List[str]) -> List[float]:
        if not sentences:
            return []
        
        tfidf = TfidfVectorizer(stop_words='english', max_features=100)
        try:
            tfidf_matrix = tfidf.fit_transform(sentences)
            
            sentence_scores = []
            for i in range(len(sentences)):
                scores = []
                for j in range(len(sentences)):
                    if i != j:
                        sim = cosine_similarity(tfidf_matrix[i], tfidf_matrix[j])[0][0]
                        scores.append(sim)
                sentence_scores.append(sum(scores) / max(1, len(scores)))
            
            return sentence_scores
        except:
            return [1.0] * len(sentences)
    
    def extract_important_sentences(self, text: str, top_n: int = 20) -> List[str]:
        sentences = sent_tokenize(text)
        if len(sentences) <= top_n:
            return sentences
        
        scores = self.get_sentence_importance(sentences)
        sentence_score_pairs = list(zip(sentences, scores))
        sentence_score_pairs.sort(key=lambda x: x[1], reverse=True)
        
        return [sentence for sentence, _ in sentence_score_pairs[:top_n]]

class CurriculumPlanner:
    def __init__(self):
        self.topic_hierarchy = {}
        self.difficulty_levels = ['Basic', 'Intermediate', 'Advanced']
    
    def analyze_topic_coverage(self, chunks: List[str]) -> Dict[str, Any]:
        all_text = " ".join(chunks)
        topics = self._extract_topics(all_text)
        
        topic_analysis = {}
        for topic in topics:
            topic_sentences = [s for s in sent_tokenize(all_text) if topic.lower() in s.lower()]
            topic_analysis[topic] = {
                'frequency': len(topic_sentences),
                'sentences': topic_sentences[:5],
                'difficulty': self._estimate_difficulty(topic_sentences)
            }
        
        return topic_analysis
    
    def _extract_topics(self, text: str) -> List[str]:
        sentences = sent_tokenize(text)
        topics = []
        
        pattern = r'(?:topic|section|chapter|module|lesson|unit)\s+(\d+|[a-zA-Z]+)[:\s]+([^.!?]+)'
        matches = re.findall(pattern, text, re.IGNORECASE)
        for match in matches:
            topics.append(match[1].strip())
        
        if len(topics) < 3:
            doc = self.nlp(text[:2000])
            for chunk in doc.noun_chunks:
                if len(chunk.text.split()) > 1 and len(chunk.text) < 50:
                    topics.append(chunk.text)
        
        return list(set(topics))[:10]
    
    def _estimate_difficulty(self, sentences: List[str]) -> str:
        if not sentences:
            return 'Basic'
        
        avg_length = sum(len(s.split()) for s in sentences) / max(1, len(sentences))
        avg_word_length = sum(len(word) for s in sentences for word in s.split()) / max(1, sum(len(s.split()) for s in sentences))
        
        unique_ratio = len(set(" ".join(sentences).split())) / max(1, len(" ".join(sentences).split()))
        
        if avg_length < 12 and avg_word_length < 5 and unique_ratio < 0.5:
            return 'Basic'
        elif avg_length < 18 and avg_word_length < 7 and unique_ratio < 0.7:
            return 'Intermediate'
        else:
            return 'Advanced'

class QuestionTypeManager:
    def __init__(self):
        self.question_templates = {
            'mcq': {
                'formats': ['standard', 'true_false', 'matching', 'fill_blank'],
                'default': 'standard'
            },
            'short': {
                'formats': ['definition', 'explanation', 'comparison', 'example'],
                'default': 'explanation'
            },
            'long': {
                'formats': ['essay', 'case_study', 'scenario', 'analysis'],
                'default': 'essay'
            },
            'fill_blank': {
                'formats': ['standard'],
                'default': 'standard'
            },
            'matching': {
                'formats': ['standard'],
                'default': 'standard'
            }
        }
    
    def generate_mixed_questions(self, context: str, num_questions: int = 10) -> List[Dict[str, Any]]:
        question_types = ['mcq', 'short', 'long', 'fill_blank', 'matching']
        weights = [0.4, 0.3, 0.15, 0.075, 0.075]
        
        selected_types = np.random.choice(question_types, size=num_questions, p=weights)
        
        questions = []
        for q_type in selected_types:
            if q_type == 'mcq':
                questions.extend(self._generate_mcq_variants(context, 1))
            elif q_type == 'short':
                questions.extend(self._generate_short_variants(context, 1))
            elif q_type == 'long':
                questions.extend(self._generate_long_variants(context, 1))
            elif q_type == 'fill_blank':
                questions.extend(self._generate_fill_blank(context, 1))
            elif q_type == 'matching':
                questions.extend(self._generate_matching(context, 1))
        
        return questions[:num_questions]
    
    def _generate_mcq_variants(self, context: str, num: int) -> List[Dict[str, Any]]:
        questions = []
        sentences = sent_tokenize(context)
        
        for _ in range(num):
            if not sentences:
                break
            
            sentence = random.choice(sentences)
            if len(sentence.split()) < 5:
                continue
            
            variant = random.choice(['standard', 'true_false'])
            
            if variant == 'true_false':
                question = {
                    'question': f"True or False: {sentence}",
                    'type': 'mcq',
                    'subtype': 'true_false',
                    'options': ['True', 'False'],
                    'answer': random.choice(['True', 'False'])
                }
            else:
                options = self._generate_mcq_options(context, sentence, 4)
                question = {
                    'question': f"Choose the best answer: {sentence[:100]}...",
                    'type': 'mcq',
                    'subtype': 'standard',
                    'options': options,
                    'answer': sentence
                }
            
            questions.append(question)
        
        return questions
    
    def _generate_short_variants(self, context: str, num: int) -> List[Dict[str, Any]]:
        questions = []
        sentences = sent_tokenize(context)
        
        for _ in range(num):
            if not sentences:
                break
            
            sentence = random.choice(sentences)
            if len(sentence.split()) < 5:
                continue
            
            variant = random.choice(['definition', 'explanation', 'comparison', 'example'])
            
            templates = {
                'definition': f"Define the following term based on the material: {sentence[:50]}...",
                'explanation': f"Explain the following statement: {sentence}",
                'comparison': f"Compare and contrast the concepts related to: {sentence[:50]}...",
                'example': f"Provide an example that illustrates: {sentence[:50]}..."
            }
            
            question = {
                'question': templates.get(variant, templates['explanation']),
                'type': 'short',
                'subtype': variant,
                'answer': sentence,
                'expected_length': 50
            }
            questions.append(question)
        
        return questions
    
    def _generate_long_variants(self, context: str, num: int) -> List[Dict[str, Any]]:
        questions = []
        sentences = sent_tokenize(context)
        
        for _ in range(num):
            if len(sentences) < 5:
                break
            
            chunk = random.sample(sentences, min(5, len(sentences)))
            chunk_text = " ".join(chunk)
            
            variant = random.choice(['essay', 'case_study', 'scenario', 'analysis'])
            
            templates = {
                'essay': f"Write a comprehensive essay on: {chunk_text[:100]}...",
                'case_study': f"Analyze the following case study: {chunk_text[:100]}...",
                'scenario': f"Given the scenario: {chunk_text[:100]}... What would you conclude?",
                'analysis': f"Provide a detailed analysis of: {chunk_text[:100]}..."
            }
            
            question = {
                'question': templates.get(variant, templates['essay']),
                'type': 'long',
                'subtype': variant,
                'answer': chunk_text,
                'expected_length': 200
            }
            questions.append(question)
        
        return questions
    
    def _generate_fill_blank(self, context: str, num: int) -> List[Dict[str, Any]]:
        questions = []
        sentences = sent_tokenize(context)
        
        for _ in range(num):
            if not sentences:
                break
            
            sentence = random.choice(sentences)
            words = sentence.split()
            
            if len(words) < 5:
                continue
            
            blank_pos = random.randint(1, len(words) - 2)
            answer = words[blank_pos]
            words[blank_pos] = "_______"
            
            question = {
                'question': " ".join(words),
                'type': 'fill_blank',
                'subtype': 'standard',
                'answer': answer,
                'expected_length': 1
            }
            questions.append(question)
        
        return questions
    
    def _generate_matching(self, context: str, num: int) -> List[Dict[str, Any]]:
        questions = []
        
        sentences = sent_tokenize(context)
        if len(sentences) < 4:
            return questions
        
        pairs = []
        for _ in range(min(3, len(sentences) // 2)):
            if len(sentences) < 2:
                break
            
            left = random.choice(sentences)[:30]
            sentences.remove(left)
            right = random.choice(sentences)[:30]
            sentences.remove(right)
            pairs.append((left, right))
        
        if pairs:
            left_items = [item[0] for item in pairs]
            right_items = [item[1] for item in pairs]
            random.shuffle(right_items)
            
            matching_items = list(zip(left_items, right_items))
            
            question = {
                'question': "Match the following items:",
                'type': 'matching',
                'subtype': 'standard',
                'pairs': matching_items,
                'answer': pairs
            }
            questions.append(question)
        
        return questions
    
    def _generate_mcq_options(self, context: str, correct: str, num_options: int) -> List[str]:
        sentences = [s for s in sent_tokenize(context) if s != correct]
        wrong_answers = random.sample(sentences, min(num_options - 1, len(sentences)))
        
        options = [correct]
        options.extend(wrong_answers)
        
        while len(options) < num_options:
            filler = f"Alternative concept about {random.choice(list(set(options)))[:20]}"
            if filler not in options:
                options.append(filler)
        
        random.shuffle(options)
        return options

class ScoringAndFeedback:
    def __init__(self):
        self.rubric_templates = {
            'mcq': {
                'criteria': ['Correctness', 'Clarity', 'Relevance'],
                'max_score': 10
            },
            'short': {
                'criteria': ['Completeness', 'Accuracy', 'Clarity'],
                'max_score': 20
            },
            'long': {
                'criteria': ['Depth', 'Analysis', 'Organization', 'Evidence'],
                'max_score': 40
            }
        }
    
    def generate_rubric(self, question_type: str) -> Dict[str, Any]:
        template = self.rubric_templates.get(question_type, self.rubric_templates['short'])
        
        rubric = {
            'type': question_type,
            'criteria': template['criteria'],
            'max_score': template['max_score'],
            'scoring_table': self._create_scoring_table(template['criteria']),
            'descriptors': self._generate_descriptors(template['criteria'])
        }
        
        return rubric
    
    def _create_scoring_table(self, criteria: List[str]) -> Dict[str, Dict[str, int]]:
        scoring = {}
        for criterion in criteria:
            scoring[criterion] = {
                'Excellent': 4,
                'Good': 3,
                'Satisfactory': 2,
                'Needs Improvement': 1,
                'Inadequate': 0
            }
        return scoring
    
    def _generate_descriptors(self, criteria: List[str]) -> Dict[str, Dict[str, str]]:
        descriptors = {}
        for criterion in criteria:
            descriptors[criterion] = {
                'Excellent': f'Shows outstanding {criterion.lower()} with exceptional detail',
                'Good': f'Demonstrates good {criterion.lower()} with sufficient detail',
                'Satisfactory': f'Shows adequate {criterion.lower()} but could be improved',
                'Needs Improvement': f'Limited {criterion.lower()} that needs significant improvement',
                'Inadequate': f'Fails to demonstrate {criterion.lower()}'
            }
        return descriptors
    
    def estimate_difficulty_score(self, question: Dict[str, Any]) -> float:
        if question.get('type') == 'mcq':
            return 0.3 + random.random() * 0.3
        elif question.get('type') == 'short':
            return 0.4 + random.random() * 0.3
        elif question.get('type') == 'long':
            return 0.6 + random.random() * 0.3
        else:
            return 0.5
    
    def generate_feedback_template(self, question: Dict[str, Any], score: float) -> str:
        if question.get('type') == 'mcq':
            templates = [
                f"The answer to '{question['question'][:50]}...' is {'correct' if score > 0.7 else 'incorrect'}.",
                f"Based on your response, you {'correctly identified' if score > 0.7 else 'need to review'} the concept.",
            ]
        elif question.get('type') == 'short':
            templates = [
                f"Your response to the short question shows {'good understanding' if score > 0.7 else 'partial understanding'}.",
                f"The answer could benefit from {'more detail' if score > 0.5 else 'significant improvement'}."
            ]
        else:
            templates = [
                f"Your essay demonstrates {'strong' if score > 0.7 else 'developing'} analytical skills.",
                f"Consider {'deepening your analysis' if score > 0.5 else 'reviewing the fundamental concepts'}."
            ]
        
        return random.choice(templates)

class ExamTemplateManager:
    def __init__(self):
        self.templates = {
            'standard': {
                'title': 'Standard Exam',
                'sections': ['mcq', 'short', 'long'],
                'weights': [0.4, 0.3, 0.3]
            },
            'comprehensive': {
                'title': 'Comprehensive Assessment',
                'sections': ['mcq', 'short', 'long', 'fill_blank', 'matching'],
                'weights': [0.3, 0.25, 0.2, 0.125, 0.125]
            },
            'quick': {
                'title': 'Quick Assessment',
                'sections': ['mcq', 'short'],
                'weights': [0.6, 0.4]
            },
            'in_depth': {
                'title': 'In-Depth Analysis',
                'sections': ['short', 'long'],
                'weights': [0.4, 0.6]
            }
        }
    
    def get_template(self, template_name: str) -> Dict[str, Any]:
        return self.templates.get(template_name, self.templates['standard'])
    
    def create_custom_template(self, sections: List[str], weights: List[float], 
                              title: str = 'Custom Exam') -> Dict[str, Any]:
        return {
            'title': title,
            'sections': sections,
            'weights': weights
        }
    
    def recommend_template(self, num_questions: int, complexity: str) -> str:
        if num_questions <= 10:
            if complexity == 'easy':
                return 'quick'
            else:
                return 'standard'
        elif num_questions <= 20:
            return 'comprehensive'
        else:
            return 'in_depth'

class AdvancedExamPaperGenerator:
    def __init__(self):
        self.template_manager = ExamTemplateManager()
        self.scoring = ScoringAndFeedback()
        self.question_manager = QuestionTypeManager()
    
    def generate_paper_with_template(self, questions: List[Dict[str, Any]], 
                                    template_name: str = 'standard',
                                    title: str = "Exam Paper") -> BytesIO:
        template = self.template_manager.get_template(template_name)
        
        organized_questions = []
        for section in template['sections']:
            section_questions = [q for q in questions if q.get('type') == section]
            organized_questions.extend(section_questions)
        
        return self.generate_word_advanced(organized_questions, title)
    
    def generate_word_advanced(self, questions: List[Dict[str, Any]], 
                               title: str = "Exam Paper") -> BytesIO:
        buffer = BytesIO()
        doc = docx.Document()
        
        doc_metadata = doc.core_properties
        doc_metadata.title = title
        doc_metadata.created = datetime.now()
        doc_metadata.creator = "Exam Paper Generator"
        
        title_paragraph = doc.add_heading(title, 0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        info_table = doc.add_table(rows=2, cols=2)
        info_table.style = 'Table Grid'
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        cells = [
            ('Generated Date', datetime.now().strftime('%B %d, %Y')),
            ('Total Questions', str(len(questions))),
            ('Instructions', 'Answer all questions to the best of your ability.'),
            ('Time Allowed', '2 hours')
        ]
        
        for i, (label, value) in enumerate(cells):
            row = i // 2
            col = i % 2
            cell = info_table.cell(row, col)
            cell.text = f"{label}: {value}"
        
        doc.add_page_break()
        
        section_counter = {
            'mcq': 0,
            'short': 0,
            'long': 0,
            'fill_blank': 0,
            'matching': 0
        }
        
        for q in questions:
            q_type = q.get('type', 'short')
            section_counter[q_type] = section_counter.get(q_type, 0) + 1
        
        section_titles = {
            'mcq': 'Section A: Multiple Choice Questions',
            'short': 'Section B: Short Answer Questions',
            'long': 'Section C: Long Answer Questions',
            'fill_blank': 'Section D: Fill in the Blanks',
            'matching': 'Section E: Matching Questions'
        }
        
        current_section = None
        section_number = 1
        
        for q_type in ['mcq', 'short', 'long', 'fill_blank', 'matching']:
            type_questions = [q for q in questions if q.get('type') == q_type]
            
            if not type_questions:
                continue
            
            if current_section:
                doc.add_page_break()
            
            doc.add_heading(f"{section_titles.get(q_type, q_type.upper())}", 1)
            doc.add_paragraph(f"Total Questions: {len(type_questions)}")
            doc.add_paragraph("Instructions: Answer each question carefully.")
            doc.add_paragraph()
            
            rubric = self.scoring.generate_rubric(q_type)
            rubric_table = doc.add_table(rows=1, cols=len(rubric['criteria']) + 1)
            rubric_table.style = 'Table Grid'
            
            header_cells = rubric_table.rows[0].cells
            header_cells[0].text = "Criteria"
            for idx, criterion in enumerate(rubric['criteria'], 1):
                header_cells[idx].text = criterion
            
            for idx, q in enumerate(type_questions, 1):
                p = doc.add_paragraph()
                run = p.add_run(f"{idx}. {q['question']}")
                run.bold = True
                run.font.size = Pt(12)
                
                if q_type == 'mcq':
                    options = q.get('options', [])
                    for opt_idx, option in enumerate(options, 1):
                        p = doc.add_paragraph()
                        p.add_run(f"    {chr(64 + opt_idx)}. {option}")
                        p.paragraph_format.left_indent = Inches(0.3)
                    
                    doc.add_paragraph(f"[Score: {rubric['max_score']} points]")
                    
                elif q_type == 'short':
                    doc.add_paragraph(f"Expected length: {q.get('expected_length', 50)} words")
                    doc.add_paragraph("Answer:")
                    doc.add_paragraph("_" * 80)
                    doc.add_paragraph()
                    
                    doc.add_paragraph(f"[Score: {rubric['max_score']} points]")
                    
                elif q_type == 'long':
                    doc.add_paragraph(f"Expected length: {q.get('expected_length', 200)} words")
                    doc.add_paragraph("Answer:")
                    doc.add_paragraph("_" * 80)
                    doc.add_paragraph()
                    doc.add_paragraph("_" * 80)
                    
                    doc.add_paragraph(f"[Score: {rubric['max_score']} points]")
                
                elif q_type == 'fill_blank':
                    doc.add_paragraph("Fill in the blank with the correct word:")
                    doc.add_paragraph()
                    
                    doc.add_paragraph(f"[Score: {rubric['max_score']} points]")
                
                elif q_type == 'matching':
                    pairs = q.get('pairs', [])
                    if pairs:
                        doc.add_paragraph("Match the following:")
                        matching_table = doc.add_table(rows=len(pairs), cols=2)
                        matching_table.style = 'Table Grid'
                        
                        for row_idx, (left, right) in enumerate(pairs):
                            matching_table.cell(row_idx, 0).text = f"{chr(65 + row_idx)}. {left}"
                            matching_table.cell(row_idx, 1).text = f"{chr(65 + row_idx)}. {right}"
                    
                    doc.add_paragraph(f"[Score: {rubric['max_score']} points]")
                
                doc.add_paragraph()
        
        doc.add_page_break()
        doc.add_heading('Answer Key', 1)
        
        for idx, q in enumerate(questions, 1):
            p = doc.add_paragraph()
            run = p.add_run(f"{idx}. ")
            run.bold = True
            
            if q.get('type') == 'mcq':
                answer_text = q.get('answer', '')
                if q.get('subtype') == 'true_false':
                    answer_text = q.get('answer', 'True')
                p.add_run(f"Answer: {answer_text[:100]}")
            elif q.get('type') in ['short', 'long']:
                p.add_run(f"Answer: {q.get('answer', '')[:150]}")
            elif q.get('type') == 'fill_blank':
                p.add_run(f"Answer: {q.get('answer', '')}")
            elif q.get('type') == 'matching':
                pairs = q.get('answer', [])
                if pairs:
                    p.add_run(f"Answers: {pairs}")
        
        doc.save(buffer)
        buffer.seek(0)
        return buffer

class EnhancedExamApp(ExamPaperApp):
    def __init__(self):
        super().__init__()
        self.advanced_processor = AdvancedTextProcessor()
        self.curriculum_planner = CurriculumPlanner()
        self.advanced_generator = AdvancedExamPaperGenerator()
        self.template_manager = ExamTemplateManager()
        self.session_state = st.session_state
        
        if 'history' not in self.session_state:
            self.session_state.history = []
        if 'favorites' not in self.session_state:
            self.session_state.favorites = []
    
    def run(self):
        st.set_page_config(
            page_title="Advanced Exam Paper Generator",
            page_icon="📚",
            layout="wide",
            initial_sidebar_state="expanded"
        )
        
        st.markdown("""
            <style>
                .main-header { font-size: 2.5rem; color: #1a1a2e; text-align: center; }
                .sub-header { font-size: 1.2rem; color: #4a4a6a; text-align: center; margin-bottom: 2rem; }
                .stButton > button { width: 100%; }
                .uploaded-file { background-color: #f0f2f6; padding: 0.5rem; border-radius: 0.5rem; margin: 0.2rem 0; }
                .question-preview { background-color: #f8f9fa; padding: 1rem; border-radius: 0.5rem; margin: 0.5rem 0; }
                .stTabs [data-baseweb="tab-list"] { gap: 2rem; }
                .stTabs [data-baseweb="tab"] { height: 3rem; white-space: pre-wrap; }
            </style>
        """, unsafe_allow_html=True)
        
        st.markdown('<p class="main-header">📚 Advanced Exam Paper Generator</p>', unsafe_allow_html=True)
        st.markdown('<p class="sub-header">RAG-powered application for intelligent exam paper creation</p>', unsafe_allow_html=True)
        
        tabs = st.tabs(["📝 Generate Exam", "📊 Analyze Content", "📁 Template Manager", "📈 History & Favorites"])
        
        with tabs[0]:
            self._render_generation_tab()
        
        with tabs[1]:
            self._render_analysis_tab()
        
        with tabs[2]:
            self._render_template_tab()
        
        with tabs[3]:
            self._render_history_tab()
    
    def _render_generation_tab(self):
        col1, col2 = st.columns([1, 2])
        
        with col1:
            st.subheader("📚 Material Management")
            
            uploaded_files = st.file_uploader(
                "Upload teaching materials",
                type=['pdf', 'docx', 'txt', 'md'],
                accept_multiple_files=True,
                key="file_uploader_main"
            )
            
            if uploaded_files:
                if st.button("🔄 Process Documents", type="primary", use_container_width=True):
                    with st.spinner("Processing documents..."):
                        for file in uploaded_files:
                            try:
                                text = self.document_processor.extract_text(
                                    file.getvalue(), 
                                    file.name
                                )
                                chunks = self.document_processor.chunk_text(text, chunk_size=500)
                                
                                metadata = {
                                    'filename': file.name,
                                    'upload_time': datetime.now().isoformat(),
                                    'chunks': len(chunks)
                                }
                                
                                self.vector_store.add_documents(chunks, metadata)
                                
                                self.uploaded_files.append({
                                    'name': file.name,
                                    'chunks': len(chunks)
                                })
                            except Exception as e:
                                st.error(f"Error processing {file.name}: {str(e)}")
                        
                        st.success(f"✅ Processed {len(uploaded_files)} files successfully!")
                        st.rerun()
            
            if self.uploaded_files:
                st.subheader("📄 Uploaded Documents")
                for file_info in self.uploaded_files:
                    st.markdown(f"- ✅ {file_info['name']} ({file_info['chunks']} chunks)")
                
                if st.button("🗑️ Clear All Materials", use_container_width=True):
                    self.vector_store.clear()
                    self.uploaded_files = []
                    st.success("All materials cleared!")
                    st.rerun()
            
            st.subheader("⚙️ Generation Settings")
            
            template_options = list(self.template_manager.templates.keys())
            selected_template = st.selectbox("Exam Template", template_options)
            
            num_questions_total = st.slider(
                "Total Questions",
                min_value=5,
                max_value=50,
                value=20,
                step=5
            )
            
            question_types = st.multiselect(
                "Question Types",
                ['MCQ', 'Short Answer', 'Long Answer', 'Fill in the Blank', 'Matching'],
                default=['MCQ', 'Short Answer']
            )
            
            difficulty = st.select_slider(
                "Difficulty Level",
                options=['Easy', 'Medium', 'Hard'],
                value='Medium'
            )
            
            include_answer_key = st.checkbox("Include Answer Key", value=True)
            
            output_format = st.radio(
                "Output Format",
                ['PDF', 'Word'],
                index=1
            )
            
            if st.button("🚀 Generate Exam Paper", type="primary", use_container_width=True):
                self._generate_enhanced_paper(
                    num_questions_total,
                    question_types,
                    difficulty,
                    include_answer_key,
                    output_format,
                    selected_template
                )
        
        with col2:
            st.subheader("📖 Content Preview")
            
            if self.vector_store.collection.count() > 0:
                st.success(f"📊 Total chunks in database: {self.vector_store.collection.count()}")
                
                search_query = st.text_input("🔍 Search your material", placeholder="Enter a topic or keyword...")
                
                if search_query:
                    results = self.vector_store.query(search_query, n_results=3)
                    if results:
                        st.markdown("**Top matches:**")
                        for i, result in enumerate(results, 1):
                            with st.expander(f"Result {i}"):
                                st.write(result[:300] + "...")
                
                if st.button("📊 Analyze Material", use_container_width=True):
                    with st.spinner("Analyzing content..."):
                        self._show_content_analysis()
            else:
                st.info("📭 No materials uploaded. Please upload teaching material to begin.")
                st.markdown("""
                    ### 💡 Quick Start
                    1. Upload your teaching materials (PDF, DOCX, TXT)
                    2. Configure your exam settings
                    3. Generate a customized exam paper
                """)
    
    def _show_content_analysis(self):
        with st.spinner("Analyzing material..."):
            all_chunks = self.vector_store.collection.get()['documents']
            
            if not all_chunks:
                st.warning("No content to analyze")
                return
            
            full_text = " ".join(all_chunks)
            
            keywords = self.advanced_processor.extract_keywords(full_text, top_n=20)
            key_phrases = self.advanced_processor.extract_key_phrases(full_text, top_n=10)
            important_sentences = self.advanced_processor.extract_important_sentences(full_text, top_n=10)
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("#### 🔑 Key Topics")
                for phrase in key_phrases[:5]:
                    st.markdown(f"- {phrase.title()}")
            
            with col2:
                st.markdown("#### 📊 Content Statistics")
                st.metric("Total Words", len(full_text.split()))
                st.metric("Total Sentences", len(sent_tokenize(full_text)))
                st.metric("Unique Words", len(set(full_text.split())))
            
            st.markdown("#### 📝 Key Sentences")
            for i, sentence in enumerate(important_sentences[:5], 1):
                st.markdown(f"{i}. {sentence}")
    
    def _render_analysis_tab(self):
        st.subheader("📊 Content Analysis Dashboard")
        
        if self.vector_store.collection.count() == 0:
            st.info("Upload materials first to see analysis")
            return
        
        all_chunks = self.vector_store.collection.get()['documents']
        full_text = " ".join(all_chunks)
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("Total Documents", len(self.uploaded_files))
            st.metric("Total Chunks", len(all_chunks))
        
        with col2:
            sentences = sent_tokenize(full_text)